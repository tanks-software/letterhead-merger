import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import base64
import io
import os
import subprocess
import time
from PIL import Image
from streamlit_cropper import st_cropper
from docx.shared import Inches
from drive_utils import list_files_in_folder, download_file
from tempfile import NamedTemporaryFile
from pathlib import Path

st.set_page_config(page_title="Letterhead Merger", layout="wide")
st.title("📄 Letterhead Document Merger (Drive-Based)")

# === Google Drive Folder IDs ===
LETTERHEAD_FOLDER_ID = "1cptQfvNP9UxHK_-lfkZc6lEqW3GfAg4e"
BODY_FOLDER_ID = "1d87BF8jSmyTibx3-qYgP_gG4Z9mLsFTr"

# === Select Letterhead from Drive ===
letterhead_files = list_files_in_folder(LETTERHEAD_FOLDER_ID, ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"])
selected_letterhead = st.selectbox("📁 Select LETTERHEAD File from Google Drive", [f['name'] for f in letterhead_files])
letterhead_file_id = next(f['id'] for f in letterhead_files if f['name'] == selected_letterhead)
letterhead_bytes = download_file(letterhead_file_id)

# === Select Body File from Drive ===
body_files = list_files_in_folder(BODY_FOLDER_ID, ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"])
selected_body = st.selectbox("📁 Select BODY File from Google Drive", [f['name'] for f in body_files])
body_file_id = next(f['id'] for f in body_files if f['name'] == selected_body)
body_bytes = download_file(body_file_id)

# === Optional Reference PDF Upload ===
st.subheader("📎 Upload Reference PDF (optional)")
ref_pdf = st.file_uploader("Upload PDF to preview", type=["pdf"])
if ref_pdf:
    base64_pdf = base64.b64encode(ref_pdf.read()).decode("utf-8")
    st.markdown(f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600"></iframe>', unsafe_allow_html=True)

# === Helper Functions ===
def docx_to_text_from_bytes(byte_stream):
    doc = Document(byte_stream)
    return "\n".join(p.text for p in doc.paragraphs)

def build_clean_letterhead_docx(header_img, footer_img, body_text, letterhead_signature_lines, signature_image=None):
    doc = Document()
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    if section.header.paragraphs:
        section.header.paragraphs[0].add_run().add_picture(str(header_img), width=Inches(6))

    for line in body_text.splitlines():
        doc.add_paragraph(line)

    for line in letterhead_signature_lines:
        doc.add_paragraph(line)

    if signature_image and Path(signature_image).exists():
        with open(signature_image, "rb") as f:
            doc.add_picture(f, width=Inches(2))

    if section.footer.paragraphs:
        section.footer.paragraphs[0].add_run().add_picture(str(footer_img), width=Inches(6))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === Process Letterhead File ===
signature_crop_path = None
pdf_img_path = "letterhead_preview.png"

if selected_letterhead.endswith(".docx"):
    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(letterhead_bytes.read())
        temp_docx.flush()
        pdf_name = Path(temp_docx.name).with_suffix(".pdf").name
        output_pdf_path = os.path.join(os.getcwd(), pdf_name)
        result = subprocess.run(["soffice", "--headless", "--convert-to", "pdf", temp_docx.name, "--outdir", os.getcwd()],
                                stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        time.sleep(2)
        if not os.path.exists(output_pdf_path):
            st.error(f"❌ Failed to convert DOCX to PDF.\n\nCommand: soffice ...\nExit Code: {result.returncode}\n\nSTDOUT:\n{result.stdout.decode()}\n\nSTDERR:\n{result.stderr.decode()}")
            st.stop()
        temp_pdf_path = output_pdf_path
else:
    with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf.write(letterhead_bytes.read())
        temp_pdf.flush()
        temp_pdf_path = temp_pdf.name

# ✅ Use PyMuPDF instead of pdf2image
doc = fitz.open(temp_pdf_path)
page = doc.load_page(0)
pix = page.get_pixmap(dpi=150)
pix.save(pdf_img_path)
image = Image.open(pdf_img_path)

st.subheader("✂️ Crop Signature From Letterhead Preview")
cropped_img = st_cropper(image, box_color='#0000FF', aspect_ratio=None)
cropped_img_path = "cropped_signature.png"
cropped_img.save(cropped_img_path)
st.image(cropped_img_path, caption="Selected Signature Area", width=300)

# === Body Text Edit ===
st.subheader("📝 Edit Body Text")
body_text = docx_to_text_from_bytes(body_bytes)
edited_text = st.text_area("Edit Body Content", body_text, height=400)

# === Generate Final DOCX ===
if st.button("📥 Generate & Download Final DOCX"):
    header_img = "temp_header.png"
    footer_img = "temp_footer.png"

    page.get_pixmap(clip=fitz.Rect(0, 0, page.rect.width, 100)).save(header_img)
    page.get_pixmap(clip=fitz.Rect(0, page.rect.height - 100, page.rect.width, page.rect.height)).save(footer_img)

    final_docx = build_clean_letterhead_docx(header_img, footer_img, edited_text, [], cropped_img_path)

    letterhead_name = selected_letterhead.replace(" ", "_").replace(".docx", "").replace(".pdf", "")
    body_name = selected_body.replace(" ", "_").replace(".docx", "")
    final_filename = f"letterhead({letterhead_name})_blsurrender({body_name}).docx"

    st.success(f"✅ Ready to Download: {final_filename}")
    st.download_button(
        label="📥 Download DOCX",
        data=final_docx,
        file_name=final_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
